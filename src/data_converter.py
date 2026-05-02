from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO, StringIO
import re
from typing import Any

import pandas as pd


TEMPLATE_COLUMNS = ["date", "country", "market", "commodity", "price", "currency", "unit"]

KEYWORDS = {
    "date": ["date", "month", "period", "time", "year"],
    "country": ["country", "nation"],
    "market": ["market", "admin", "province", "district", "location", "area", "site"],
    "commodity": ["commodity", "product", "item", "food", "staple", "crop"],
    "price": ["price", "cost", "value", "amount", "retail", "wholesale", "median"],
    "currency": ["currency", "curr", "iso"],
    "unit": ["unit", "uom", "measure", "kg", "liter", "litre"],
}


@dataclass
class ConversionResult:
    converted: pd.DataFrame
    source_preview: pd.DataFrame
    mapping: dict[str, str]
    messages: list[str]


def convert_uploaded_file(
    uploaded_file: Any,
    default_country: str = "",
    default_market: str = "",
    default_currency: str = "USD",
    default_unit: str = "kg",
) -> ConversionResult:
    """Extract tabular data from common files and convert it to the app CSV template."""
    tables, messages = extract_tables(uploaded_file)
    if not tables:
        raise ValueError("No readable table was found in the uploaded file.")

    best_table = max(tables, key=lambda table: _template_score(table))
    converted, mapping, conversion_messages = convert_table_to_template(
        best_table,
        default_country=default_country,
        default_market=default_market,
        default_currency=default_currency,
        default_unit=default_unit,
    )
    messages.extend(conversion_messages)
    return ConversionResult(
        converted=converted,
        source_preview=best_table.head(50),
        mapping=mapping,
        messages=messages,
    )


def extract_tables(uploaded_file: Any) -> tuple[list[pd.DataFrame], list[str]]:
    """Read CSV, Excel, Word, or PDF files into one or more candidate tables."""
    file_name = uploaded_file.name.lower()
    content = uploaded_file.getvalue()
    messages: list[str] = []

    if file_name.endswith(".csv"):
        return [pd.read_csv(BytesIO(content))], ["Read CSV file directly."]

    if file_name.endswith((".xlsx", ".xls")):
        sheets = pd.read_excel(BytesIO(content), sheet_name=None)
        tables = [table for table in sheets.values() if not table.empty]
        return tables, [f"Read {len(tables)} Excel sheet(s)."]

    if file_name.endswith(".docx"):
        tables = _extract_docx_tables(content)
        return tables, [f"Extracted {len(tables)} table(s) from Word document."]

    if file_name.endswith(".pdf"):
        tables = _extract_pdf_tables(content)
        return tables, [f"Extracted {len(tables)} candidate table(s) from PDF text."]

    raise ValueError("Unsupported file type. Upload CSV, XLS, XLSX, DOCX, or PDF.")


def convert_table_to_template(
    table: pd.DataFrame,
    default_country: str = "",
    default_market: str = "",
    default_currency: str = "USD",
    default_unit: str = "kg",
) -> tuple[pd.DataFrame, dict[str, str], list[str]]:
    """Map arbitrary table columns into the required upload template."""
    clean_table = table.copy()
    clean_table.columns = [str(column).strip() for column in clean_table.columns]
    clean_table = clean_table.dropna(how="all")
    mapping = infer_column_mapping(clean_table)
    messages: list[str] = []

    converted = pd.DataFrame()
    for target in TEMPLATE_COLUMNS:
        source = mapping.get(target)
        if source:
            converted[target] = clean_table[source]
        else:
            converted[target] = ""

    converted["country"] = converted["country"].replace("", pd.NA).fillna(default_country)
    converted["market"] = converted["market"].replace("", pd.NA).fillna(default_market)
    converted["currency"] = converted["currency"].replace("", pd.NA).fillna(default_currency)
    converted["unit"] = converted["unit"].replace("", pd.NA).fillna(default_unit)

    if not mapping.get("date"):
        converted["date"] = pd.date_range("2025-01-01", periods=len(converted), freq="MS")
        messages.append("No date column was detected; generated monthly dates starting 2025-01-01.")

    converted["date"] = pd.to_datetime(converted["date"], errors="coerce")
    converted["price"] = pd.to_numeric(_strip_number_text(converted["price"]), errors="coerce")

    required_text = ["country", "market", "commodity", "currency", "unit"]
    for column in required_text:
        converted[column] = converted[column].astype(str).str.strip()

    converted = converted.dropna(subset=["date", "price"])
    converted = converted[converted["price"] > 0]
    converted["date"] = converted["date"].dt.strftime("%Y-%m-%d")
    converted = converted[TEMPLATE_COLUMNS].reset_index(drop=True)

    if converted.empty:
        raise ValueError("The converter could not find valid date and price rows.")

    missing_targets = [target for target in TEMPLATE_COLUMNS if target not in mapping]
    if missing_targets:
        messages.append("Missing fields were filled from defaults or generated values: " + ", ".join(missing_targets))

    return converted, mapping, messages


def infer_column_mapping(table: pd.DataFrame) -> dict[str, str]:
    """Infer schema mapping with keyword and value-pattern scoring."""
    mapping: dict[str, str] = {}
    available_columns = list(table.columns)

    for target in TEMPLATE_COLUMNS:
        scored_columns = [(_score_column(target, column, table[column]), column) for column in available_columns]
        scored_columns = sorted(scored_columns, reverse=True)
        best_score, best_column = scored_columns[0]
        if best_score > 0:
            mapping[target] = best_column
            available_columns.remove(best_column)

    return mapping


def _score_column(target: str, column: str, values: pd.Series) -> int:
    normalized = _normalize(column)
    score = 0
    for keyword in KEYWORDS[target]:
        if normalized == keyword:
            score += 8
        elif keyword in normalized:
            score += 4

    sample = values.dropna().astype(str).head(30)
    if target == "date":
        parsed_dates = pd.to_datetime(sample, errors="coerce", format="mixed") if not sample.empty else pd.Series()
        score += int(parsed_dates.notna().mean() * 8) if not sample.empty else 0
    elif target == "price":
        numeric = pd.to_numeric(_strip_number_text(sample), errors="coerce")
        score += int(numeric.notna().mean() * 8) if not sample.empty else 0
    elif target in {"country", "market", "commodity"}:
        score += 2 if sample.nunique() > 1 else 0
    elif target in {"currency", "unit"}:
        score += 2 if sample.nunique() <= 10 else 0

    return score


def _extract_docx_tables(content: bytes) -> list[pd.DataFrame]:
    try:
        from docx import Document
    except ImportError as error:
        raise ValueError("Install python-docx to read Word files.") from error

    document = Document(BytesIO(content))
    tables = []
    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        parsed = _rows_to_dataframe(rows)
        if not parsed.empty:
            tables.append(parsed)

    if tables:
        return tables

    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    return _tables_from_text(text)


def _extract_pdf_tables(content: bytes) -> list[pd.DataFrame]:
    try:
        from pypdf import PdfReader
    except ImportError as error:
        raise ValueError("Install pypdf to read PDF files.") from error

    reader = PdfReader(BytesIO(content))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    return _tables_from_text(text)


def _tables_from_text(text: str) -> list[pd.DataFrame]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return []

    delimiter = _detect_delimiter(lines)
    rows = []
    for line in lines:
        if delimiter:
            parts = [part.strip() for part in line.split(delimiter)]
        else:
            parts = [part.strip() for part in re.split(r"\s{2,}", line)]
        if len(parts) >= 3:
            rows.append(parts)

    table = _rows_to_dataframe(rows)
    return [table] if not table.empty else []


def _rows_to_dataframe(rows: list[list[str]]) -> pd.DataFrame:
    rows = [row for row in rows if len(row) >= 3]
    if not rows:
        return pd.DataFrame()

    max_width = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (max_width - len(row)) for row in rows]
    header = normalized_rows[0]
    body = normalized_rows[1:]

    if _looks_like_header(header):
        return pd.DataFrame(body, columns=header)

    columns = [f"column_{index + 1}" for index in range(max_width)]
    return pd.DataFrame(normalized_rows, columns=columns)


def _detect_delimiter(lines: list[str]) -> str | None:
    delimiters = [",", "\t", ";", "|"]
    scored = [(sum(line.count(delimiter) for line in lines[:20]), delimiter) for delimiter in delimiters]
    count, delimiter = max(scored)
    return delimiter if count >= 3 else None


def _looks_like_header(row: list[str]) -> bool:
    joined = " ".join(row).lower()
    return any(keyword in joined for values in KEYWORDS.values() for keyword in values)


def _template_score(table: pd.DataFrame) -> int:
    mapping = infer_column_mapping(table)
    return len(mapping) * 10 + len(table)


def _normalize(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", str(value).lower()).strip("_")


def _strip_number_text(values: pd.Series) -> pd.Series:
    return values.astype(str).str.replace(r"[^0-9.\-]", "", regex=True)
